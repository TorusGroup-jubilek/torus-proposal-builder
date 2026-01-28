import os
import json
import datetime
from io import BytesIO
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Any

import streamlit as st
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pypdf import PdfReader
from openai import OpenAI

import html

CHECK = "✓"
TEMPLATE_FILE = "Torus_Template.docx"


# =========================
# DATA MODEL
# =========================
@dataclass
class ProposalInputs:
    client: str
    facility_name: str
    service_begin_date: str
    service_end_date: str
    service_addresses: List[str]
    days_per_week: int
    cleaning_times: str

    # Standard room counts
    num_offices: int
    num_conference_rooms: int
    num_break_rooms: int
    num_bathrooms: int

    # Custom room types
    custom_rooms: List[Dict[str, Any]]  # {"type": str, "count": int}

    # Consumables (optional: None means do not print)
    hand_soap: Optional[str]
    paper_towels: Optional[str]
    toilet_paper: Optional[str]

    # Cover page
    include_cover_page: bool
    cover_letter_body: str

    # Optional sections
    cleaning_plan: str
    notes: str

    # Payment (optional)
    compensation_amount: Optional[float]
    compensation_basis: str  # monthly/annual/per visit/one-time clean
    net_terms_days: Optional[int]
    late_interest_percent: Optional[float]

    # Contract sections toggles (ON by default)
    include_employee_conduct: bool
    include_on_site_storage: bool
    include_compensation_section: bool
    include_modification: bool
    include_access: bool
    include_cancellation: bool

    # Contractor signature (hardcoded per your confirmation)
    contractor_printed_name: str
    contractor_title: str


# =========================
# DEFAULT COVER LETTER
# =========================
def default_cover_letter(client_name: str) -> str:
    cn = client_name.strip() or "[Client Name]"
    return f"""Hello {cn},

I want to personally take the opportunity to say thank you for considering Torus Cleaning as an option for your commercial cleaning needs. We pride ourselves as a core value based business and seek to partner with those that align with the culture we continue to build. Our capable crews of background checked staff are constantly growing to accommodate the needs of our customers. 

As the President, I have over 20 years of project and program management in both the military and corporate settings. Believe when I tell you I am no stranger to long, busy days that carry over to the next. We strive to deliver a professional, trustworthy service that affords our customers the peace of mind to know their spaces are well maintained. 

Thank you again for the opportunity to partner with you to take your spaces beyond clean enough!

Respectfully,

Kary Jubilee - President
Torus Cleaning Services
"""


# =========================
# OPENAI (AI Analyzer)
# =========================
def get_openai_client() -> OpenAI:
    key = st.secrets.get("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("Missing OPENAI_API_KEY in Streamlit secrets.")
    return OpenAI(api_key=key)


def analyze_rfp_with_ai(text: str) -> dict:
    client = get_openai_client()

    instructions = """
Return ONLY valid JSON with this exact structure:

{
  "cleaning_plan_draft": "string",
  "scope_of_work_draft": "string",
  "schedule_rows": [
    {"task": "string", "daily": true, "weekly": false, "monthly": false}
  ],
  "clarifying_questions": ["string"]
}

Rules:
- JSON only (no markdown, no explanation)
- 12–30 realistic janitorial tasks
"""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": instructions},
            {"role": "user", "content": text[:120000]},
        ],
        response_format={"type": "json_object"},
        temperature=0.2,
    )
    return json.loads(resp.choices[0].message.content)


# =========================
# FILE TEXT EXTRACTION
# =========================
def extract_text(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    data = uploaded_file.read()

    if name.endswith(".pdf"):
        reader = PdfReader(BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in reader.pages).strip()

    if name.endswith(".docx"):
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()

    try:
        return data.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


# =========================
# WORD HELPERS
# =========================
def add_heading(doc: Document, text: str):
    p = doc.add_paragraph(text)
    (p.runs[0] if p.runs else p.add_run(text)).bold = True


def add_bullet_paragraph(doc: Document, text: str):
    # Template-safe bullet: try styles, fallback to manual bullet
    for style_name in ("List Bullet", "List Paragraph", "Bullet List"):
        try:
            doc.add_paragraph(text, style=style_name)
            return
        except KeyError:
            continue
    doc.add_paragraph(f"• {text}")


def add_cover_page(doc: Document, client: str, body: str):
    doc.add_paragraph(client)
    doc.add_paragraph("")
    doc.add_paragraph("Attn: ______________________")
    doc.add_paragraph("")
    doc.add_paragraph("Re: Janitorial Services Proposal")
    doc.add_paragraph("")
    doc.add_paragraph(f"Dear {client},")
    doc.add_paragraph("")
    doc.add_paragraph(body or "")
    doc.add_paragraph("")
    doc.add_paragraph("Respectfully,")
    doc.add_paragraph("")
    doc.add_paragraph("Kary Jubilee")
    doc.add_paragraph("President")
    doc.add_paragraph("Torus Cleaning Services")
    doc.add_page_break()


def add_scope_table(doc: Document, rows: List[tuple]):
    add_heading(doc, "SCOPE OF WORK – CLEANING SCHEDULE")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    hdr[0].text = "Task"
    hdr[1].text = "Daily"
    hdr[2].text = "Weekly"
    hdr[3].text = "Monthly"

    for task, daily, weekly, monthly in rows:
        row = table.add_row().cells
        row[0].text = str(task)
        row[1].text = CHECK if bool(daily) else ""
        row[2].text = CHECK if bool(weekly) else ""
        row[3].text = CHECK if bool(monthly) else ""

    doc.add_paragraph("")


def add_signature_blocks(doc: Document, contractor_name: str, contractor_title: str):
    doc.add_paragraph("")
    add_heading(doc, "SIGNATURES")

    # Contractor block
    doc.add_paragraph("Date: ___________________")
    doc.add_paragraph("__________________________________")
    doc.add_paragraph("Contractor Signature")
    doc.add_paragraph(f"Contractor Printed Name: {contractor_name}")
    doc.add_paragraph(f"Title: {contractor_title}")

    doc.add_paragraph("")

    # Client block (blank fields)
    doc.add_paragraph("Date: ___________________")
    doc.add_paragraph("__________________________________")
    doc.add_paragraph("Client Signature")
    doc.add_paragraph("Client Printed Name: _____________________________")
    doc.add_paragraph("Client Title: ____________________________________")


# =========================
# CONTRACT SECTIONS
# =========================
def add_employee_conduct_section(doc: Document):
    add_heading(doc, "CONDUCT OF EMPLOYEES")
    doc.add_paragraph(
        "The Contractor shall be responsible for controlling employee conduct, for assuring that its employees are not boisterous or rude, "
        "and assuring that they are not engaging in any destructive or criminal activity. The Contractor is also responsible for assuring that "
        "its employees do not disturb papers on desks, open desk drawers, cabinets, briefcases, or use Client phones, except as authorized. "
        "The Contractor and its employees shall conduct themselves in a professional manner and not read newspapers, books, or similar items while at the job site. "
        "In addition, the Contractor’s employee shall not fraternize with Client’s employees while at the job site."
    )
    doc.add_paragraph(
        "The Client reserves the right to request the removal of any of the Contractor's employees from the building at any time. "
        "Such requests will be made to the Contractor’s supervisory personnel. At no time shall the Client assume the role of the supervisor of the Contractor's personnel."
    )
    doc.add_paragraph(
        "Should the Client observe any action by the Contractor's personnel that requires correction, they shall immediately report the action to the Contractor's supervisor, "
        "who in turn shall take immediate corrective measures. In the event the Contractor's supervisor does not take immediate corrective measures, "
        "the Client shall exercise the option of requesting the removal of the offending Contractor's employee from property."
    )
    doc.add_paragraph(
        "The Client will make a written report of any occurrence of misconduct by the Contractor's employees to the Contract Administrator within twenty-four (24) hours of such an occurrence. "
        "It is agreed that any of the following actions by the Contractor's employee(s) shall be cause for removal. These include but are not limited to:"
    )
    add_bullet_paragraph(doc, "Employee in any portion of the building in which their presence is not required by the work.")
    add_bullet_paragraph(doc, "Sitting on any furniture in the office areas.")
    add_bullet_paragraph(doc, "Using any office equipment or supplies in the office areas.")
    add_bullet_paragraph(doc, "Opening any drawers, cabinets, files, etc., or reading or removing any letters, documents, etc.")
    add_bullet_paragraph(doc, "Engaging in any loud, boisterous, or un-workmanlike conduct.")
    add_bullet_paragraph(doc, "Consuming food or beverage (other than water) in any area of the building other than the kitchen.")
    doc.add_paragraph("")


def add_on_site_storage_section(doc: Document):
    add_heading(doc, "ON-SITE STORAGE")
    doc.add_paragraph(
        "The Client will supply reasonable and suitable on-site storage space for such cleaning equipment and materials as the Contractor deems necessary for the performance of the Contract."
    )
    doc.add_paragraph("")


def add_compensation_section(doc: Document, amount: float, basis: str, net_terms_days: Optional[int]):
    add_heading(doc, "COMPENSATION")

    basis_norm = (basis or "").strip().lower()
    basis_label = {
        "annual": "annual",
        "monthly": "monthly",
        "per visit": "per visit",
        "one-time clean": "one-time clean",
    }.get(basis_norm, basis_norm or "annual")

    invoice_sentence = {
        "monthly": "The Client will be invoiced monthly in arrears.",
        "annual": "The Client will be invoiced annually.",
        "per visit": "The Client will be invoiced per visit upon completion of each visit.",
        "one-time clean": "The Client will be invoiced upon completion of the Services.",
    }.get(basis_norm, "The Client will be invoiced upon completion of the Services.")

    doc.add_paragraph(
        f"The Contractor will charge a flat {basis_label} fee of ${amount:,.2f} for the Services listed within this Agreement. "
        "The Compensation includes sales tax and other applicable duties as may be required by law."
    )
    doc.add_paragraph(invoice_sentence)

    if net_terms_days is not None:
        doc.add_paragraph(f"Invoices submitted by the Contractor to the Client are due within {int(net_terms_days)} days of receipt.")
    else:
        doc.add_paragraph("Invoices submitted by the Contractor to the Client are due within ____ days of receipt.")

    doc.add_paragraph(
        "The Contractor will be reimbursed for any expenses incurred in connection with providing the Services of this Agreement."
    )
    doc.add_paragraph("")


def add_interest_section(doc: Document, late_interest_percent: float):
    add_heading(doc, "INTEREST ON LATE PAYMENTS")
    doc.add_paragraph(
        f"Interest payable on any overdue amounts under this Agreement is charged at the rate of {late_interest_percent:.2f}% (percent)."
    )
    doc.add_paragraph("")


def add_modification_section(doc: Document):
    add_heading(doc, "MODIFICATION OF AGREEMENT")
    doc.add_paragraph(
        "Any amendment or modification of this Agreement or additional obligation assumed by either Party in connection with this Agreement will only be binding "
        "if evidenced in writing signed by each Party or an authorized representative of each Party."
    )
    doc.add_paragraph("")


def add_access_section(doc: Document):
    add_heading(doc, "ACCESS")
    doc.add_paragraph(
        "The Client agrees to provide the Contractor with the necessary access to the Property and all areas of the Property as defined within the Agreement."
    )
    doc.add_paragraph("")


def add_cancellation_section(doc: Document):
    add_heading(doc, "CANCELLATION")
    doc.add_paragraph(
        "This service agreement may be terminated at any time by the Client or Contractor upon mutual agreement."
    )
    doc.add_paragraph(
        "The Client understands that the Contractor may terminate this agreement at any time if the Client fails to pay for the Services provided under this Agreement "
        "or if the Client breaches any other material provision listed in this Cleaning Services Agreement. Client agrees to pay any outstanding balances within (10) ten days of termination."
    )
    doc.add_paragraph("")


# =========================
# BUILD WORD DOC
# =========================
def build_doc(p: ProposalInputs, schedule_rows: List[tuple]) -> bytes:
    doc = Document(TEMPLATE_FILE) if os.path.exists(TEMPLATE_FILE) else Document()

    for s in doc.sections:
        s.different_first_page_header_footer = False

    if p.include_cover_page:
        client_for_letter = p.client.strip() if p.client.strip() else "[Client Name]"
        add_cover_page(doc, client_for_letter, p.cover_letter_body)

    # Title
    title = doc.add_paragraph("CLEANING SERVICE AGREEMENT")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    (title.runs[0] if title.runs else title.add_run("CLEANING SERVICE AGREEMENT")).bold = True

    # Client fields
    doc.add_paragraph(f"Client: {p.client}")
    doc.add_paragraph(f"Facility: {p.facility_name}")

    # Addresses
    doc.add_paragraph("Service Address(es):")
    for a in (p.service_addresses or []):
        a2 = (a or "").strip()
        if a2:
            add_bullet_paragraph(doc, a2)
    doc.add_paragraph("")

    # Agreement paragraphs (date blank)
    client_name = p.client.strip() if p.client.strip() else "[Client Name]"
    first_addr = (p.service_addresses[0] if p.service_addresses else "[service address]")
    doc.add_paragraph(
        f"{client_name}, (‘Client’), enters into this agreement on this date ______________ "
        f"for Torus Cleaning Services (‘Contractor’), to provide janitorial services for facility/facilities located at the following locations: {first_addr}"
    )
    doc.add_paragraph(
        f"Contractor shall provide janitorial services {p.days_per_week} per week between the hours of {p.cleaning_times} "
        f"for the facility/facilities located at {first_addr}."
    )
    doc.add_paragraph(f"The contract period is as follows {p.service_begin_date} to {p.service_end_date}.")
    doc.add_paragraph("")

    # Room counts
    add_heading(doc, "ROOM COUNTS")
    add_bullet_paragraph(doc, f"Offices: {p.num_offices}")
    add_bullet_paragraph(doc, f"Conference rooms: {p.num_conference_rooms}")
    add_bullet_paragraph(doc, f"Break rooms: {p.num_break_rooms}")
    add_bullet_paragraph(doc, f"Bathrooms: {p.num_bathrooms}")

    # Custom rooms
    extras = []
    for r in (p.custom_rooms or []):
        rt = str(r.get("type", "")).strip()
        try:
            rc = int(r.get("count", 0) or 0)
        except Exception:
            rc = 0
        if rt and rc > 0:
            extras.append((rt, rc))

    if extras:
        doc.add_paragraph("Additional room types:")
        for rt, rc in extras:
            add_bullet_paragraph(doc, f"{rt}: {rc}")

    doc.add_paragraph("")

    # Scope table
    add_scope_table(doc, schedule_rows)

    # Cleaning plan
    if (p.cleaning_plan or "").strip():
        add_heading(doc, "CLEANING PLAN")
        doc.add_paragraph(p.cleaning_plan.strip())
        doc.add_paragraph("")

    # General requirements
    add_heading(doc, "GENERAL REQUIREMENTS")
    doc.add_paragraph(
        "Contractor shall provide all labor, supervision, and personnel necessary to perform the services described in this agreement. "
        "Unless otherwise stated, Contractor shall provide all standard equipment and cleaning supplies."
    )

    consumables_lines = []
    if p.hand_soap:
        consumables_lines.append(f"Hand soap: {p.hand_soap}")
    if p.paper_towels:
        consumables_lines.append(f"Paper towels: {p.paper_towels}")
    if p.toilet_paper:
        consumables_lines.append(f"Toilet paper: {p.toilet_paper}")

    if consumables_lines:
        doc.add_paragraph("")
        doc.add_paragraph("Consumable supplies:")
        for line in consumables_lines:
            doc.add_paragraph(f"• {line}")

    doc.add_paragraph("")

    # Contract sections
    if p.include_employee_conduct:
        add_employee_conduct_section(doc)

    if p.include_on_site_storage:
        add_on_site_storage_section(doc)

    # Payment sections print only if include_compensation_section AND amount entered
    if p.include_compensation_section and (p.compensation_amount is not None):
        add_compensation_section(doc, p.compensation_amount, p.compensation_basis, p.net_terms_days)
        if p.late_interest_percent is not None:
            add_interest_section(doc, p.late_interest_percent)

    if p.include_modification:
        add_modification_section(doc)

    if p.include_access:
        add_access_section(doc)

    if p.include_cancellation:
        add_cancellation_section(doc)

    # Notes
    add_heading(doc, "NOTES")
    doc.add_paragraph((p.notes or "").strip() or "(none)")

    # Signatures
    add_signature_blocks(doc, p.contractor_printed_name, p.contractor_title)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# =========================
# PRINT PREVIEW HELPERS (HTML)
# =========================
def _esc(x: str) -> str:
    return html.escape(x or "")

def schedule_rows_to_html_table(rows: List[tuple]) -> str:
    df = pd.DataFrame(rows, columns=["Task", "Daily", "Weekly", "Monthly"]).copy()
    for col in ["Daily", "Weekly", "Monthly"]:
        df[col] = df[col].apply(lambda v: CHECK if bool(v) else "")
    return df.to_html(index=False, escape=True)

def build_print_preview_html(li: dict) -> str:
    client = _esc(li.get("client", ""))
    facility = _esc(li.get("facility", ""))
    begin = _esc(li.get("begin", ""))
    end = _esc(li.get("end", ""))
    days = li.get("days", "")
    times = _esc(li.get("times", ""))
    addresses = li.get("addresses", []) or []

    offices = li.get("offices", 0)
    conference = li.get("conference", 0)
    breaks = li.get("breaks", 0)
    baths = li.get("baths", 0)

    custom_rooms = li.get("custom_rooms", []) or []
    custom_lines = []
    for r in custom_rooms:
        rt = str(r.get("type", "")).strip()
        try:
            rc = int(r.get("count", 0) or 0)
        except Exception:
            rc = 0
        if rt and rc > 0:
            custom_lines.append(f"<li>{_esc(rt)}: {rc}</li>")

    cons = li.get("consumables", {}) or {}
    cons_lines = []
    if cons.get("hand_soap"):
        cons_lines.append(f"<li>Hand soap: {_esc(cons['hand_soap'])}</li>")
    if cons.get("paper_towels"):
        cons_lines.append(f"<li>Paper towels: {_esc(cons['paper_towels'])}</li>")
    if cons.get("toilet_paper"):
        cons_lines.append(f"<li>Toilet paper: {_esc(cons['toilet_paper'])}</li>")

    pay = li.get("payment", {}) or {}
    amount = pay.get("amount")
    basis = (pay.get("basis") or "").strip()
    net_terms = pay.get("net_terms")
    late_interest = pay.get("late_interest")

    basis_norm = basis.lower()
    invoice_sentence = {
        "monthly": "The Client will be invoiced monthly in arrears.",
        "annual": "The Client will be invoiced annually.",
        "per visit": "The Client will be invoiced per visit upon completion of each visit.",
        "one-time clean": "The Client will be invoiced upon completion of the Services.",
    }.get(basis_norm, "The Client will be invoiced upon completion of the Services.")

    schedule_rows = li.get("schedule_rows", []) or []
    schedule_table_html = schedule_rows_to_html_table(schedule_rows)

    addr_html = "".join(f"<li>{_esc(a)}</li>" for a in addresses) if addresses else "<li>(none)</li>"
    first_addr = addresses[0] if addresses else "[service address]"
    first_addr_esc = _esc(first_addr)

    cleaning_plan = _esc(li.get("cleaning_plan", "") or "")
    notes = _esc(li.get("notes", "") or "")

    secs = li.get("sections", {}) or {}
    included = []
    if secs.get("employee_conduct"): included.append("CONDUCT OF EMPLOYEES")
    if secs.get("on_site_storage"): included.append("ON-SITE STORAGE")
    if secs.get("compensation") and (amount is not None): included.append("COMPENSATION / INTEREST ON LATE PAYMENTS")
    if secs.get("modification"): included.append("MODIFICATION OF AGREEMENT")
    if secs.get("access"): included.append("ACCESS")
    if secs.get("cancellation"): included.append("CANCELLATION")
    included_html = "".join(f"<li>{_esc(s)}</li>" for s in included) if included else "<li>(none)</li>"

    payment_html = ""
    if amount is None:
        payment_html = "<p class='muted'>(Compensation section will not print — no amount entered.)</p>"
    else:
        payment_html = f"""
        <p>The Contractor will charge a flat <b>{_esc(basis)}</b> fee of <b>${amount:,.2f}</b> for the Services listed within this Agreement.</p>
        <p>{_esc(invoice_sentence)}</p>
        <p>Invoices are due within <b>{_esc(str(net_terms) if net_terms is not None else '____')}</b> days of receipt.</p>
        """
        if late_interest is not None:
            payment_html += f"<p><b>Interest on late payments:</b> {late_interest:.2f}%</p>"

    consumables_html = "<ul>" + "".join(cons_lines) + "</ul>" if cons_lines else "<p class='muted'>(No consumables will print.)</p>"
    custom_rooms_html = "<ul>" + "".join(custom_lines) + "</ul>" if custom_lines else "<p class='muted'>(No additional room types will print.)</p>"

    cleaning_plan_html = f"<p>{cleaning_plan}</p>" if cleaning_plan.strip() else "<p class='muted'>(Not included.)</p>"
    notes_html = f"<p>{notes}</p>" if notes.strip() else "<p class='muted'>(none)</p>"

    html_out = f"""
    <div class="page">
      <div class="doc-title">CLEANING SERVICE AGREEMENT</div>

      <p><b>Client:</b> {client}<br/>
         <b>Facility:</b> {facility}</p>

      <p><b>Service Address(es):</b></p>
      <ul>{addr_html}</ul>

      <p>{client if client else "[Client Name]"}, (‘Client’), enters into this agreement on this date ______________
      for Torus Cleaning Services (‘Contractor’), to provide janitorial services for facility/facilities located at the following locations:
      <b>{first_addr_esc}</b></p>

      <p>Contractor shall provide janitorial services <b>{days}</b> per week between the hours of <b>{times}</b>
      for the facility/facilities located at <b>{first_addr_esc}</b>.</p>

      <p>The contract period is as follows <b>{begin}</b> to <b>{end}</b>.</p>

      <h3>ROOM COUNTS</h3>
      <ul>
        <li>Offices: {offices}</li>
        <li>Conference rooms: {conference}</li>
        <li>Break rooms: {breaks}</li>
        <li>Bathrooms: {baths}</li>
      </ul>

      <p><b>Additional room types:</b></p>
      {custom_rooms_html}

      <h3>SCOPE OF WORK – CLEANING SCHEDULE</h3>
      <div class="table-wrap">{schedule_table_html}</div>

      <h3>CLEANING PLAN</h3>
      {cleaning_plan_html}

      <h3>GENERAL REQUIREMENTS</h3>
      <p>Contractor shall provide all labor, supervision, and personnel necessary to perform the services described in this agreement.
      Unless otherwise stated, Contractor shall provide all standard equipment and cleaning supplies.</p>

      <p><b>Consumable supplies:</b></p>
      {consumables_html}

      <h3>COMPENSATION</h3>
      {payment_html}

      <h3>INCLUDED CONTRACT SECTIONS</h3>
      <ul>{included_html}</ul>

      <h3>NOTES</h3>
      {notes_html}

      <h3>SIGNATURES</h3>
      <p><b>Date:</b> ___________________<br/>
         _________________________________<br/>
         Contractor Signature<br/>
         Contractor Printed Name: Kary Jubilee<br/>
         Title: President, Torus Cleaning Services</p>

      <p><b>Date:</b> ___________________<br/>
         _________________________________<br/>
         Client Signature<br/>
         Client Printed Name: _____________________________<br/>
         Client Title: ____________________________________</p>
    </div>
    """
    return html_out


# =========================
# STREAMLIT UI (Single-page)
# =========================
st.set_page_config(layout="wide")
st.title("Torus Group – Cleaning Proposal Builder")

# Session defaults
st.session_state.setdefault("ai", None)
st.session_state.setdefault("cover_body_custom", "")
st.session_state.setdefault("custom_rooms", [{"type": "", "count": 0}])
st.session_state.setdefault("last_inputs", None)
st.session_state.setdefault(
    "schedule_df",
    pd.DataFrame(
        [
            ("Empty trash & replace liners", True, False, False),
            ("Clean & disinfect restrooms", True, False, False),
            ("Vacuum carpet / sweep hard floors", True, False, False),
            ("Wipe high-touch points (handles, switches)", True, False, False),
            ("Dust reachable surfaces", False, True, False),
            ("Mop hard floors (as applicable)", False, True, False),
            ("Clean break room counters & sink", False, True, False),
            ("Glass/mirrors touch-up", False, True, False),
            ("High dusting (vents/ledges)", False, False, True),
            ("Detail baseboards/edges", False, False, True),
        ],
        columns=["Task", "Daily", "Weekly", "Monthly"],
    ),
)

# iPad-friendly buttons (outside the form)
top1, top2, top3, top4 = st.columns([1, 1, 1, 2])
with top1:
    if st.button("➕ Add room row"):
        st.session_state["custom_rooms"].append({"type": "", "count": 0})
        st.rerun()
with top2:
    if st.button("➖ Remove last room row"):
        if len(st.session_state["custom_rooms"]) > 1:
            st.session_state["custom_rooms"].pop()
            st.rerun()
with top3:
    if st.button("🧹 Add schedule row"):
        df = st.session_state["schedule_df"]
        df.loc[len(df)] = ["", False, False, False]
        st.session_state["schedule_df"] = df
        st.rerun()
with top4:
    st.caption(f"Template found: {os.path.exists(TEMPLATE_FILE)}  |  Template file: {TEMPLATE_FILE}")

with st.form("proposal_form", clear_on_submit=False):
    st.subheader("Client & Contract")
    c1, c2, c3 = st.columns(3)
    with c1:
        client = st.text_input("Client", value="")
        facility = st.text_input("Facility name", value="")
    with c2:
        service_begin_date = st.text_input("Service begin date", value="")
        service_end_date = st.text_input("Service end date", value="")
    with c3:
        days = st.number_input("Days per week", min_value=1, value=5)
        times = st.text_input("Cleaning times (e.g., 6 PM – 10 PM)", value="")

    st.subheader("Service Addresses")
    addresses_text = st.text_area("One address per line", height=110)

    st.subheader("Room Counts (Standard)")
    r1, r2, r3, r4 = st.columns(4)
    with r1:
        offices = st.number_input("Offices", min_value=0, value=0)
    with r2:
        conference = st.number_input("Conference rooms", min_value=0, value=0)
    with r3:
        breaks = st.number_input("Break rooms", min_value=0, value=0)
    with r4:
        baths = st.number_input("Bathrooms", min_value=0, value=0)

    st.subheader("Additional Room Types (Name + Count)")
    st.caption("Rows with blank name or 0 count will not print.")
    for i, room in enumerate(st.session_state["custom_rooms"]):
        rc1, rc2 = st.columns([3, 1])
        with rc1:
            st.session_state["custom_rooms"][i]["type"] = st.text_input(
                f"custom_room_type_{i}",
                value=str(room.get("type", "")),
                placeholder="e.g., Exam Rooms",
            )
        with rc2:
            st.session_state["custom_rooms"][i]["count"] = st.number_input(
                f"custom_room_count_{i}",
                min_value=0,
                step=1,
                value=int(room.get("count", 0) or 0),
            )

    st.subheader("Consumables (Optional)")
    st.caption("Leave blank if not included. Only selected items will appear in the agreement.")
    cA, cB, cC = st.columns(3)
    with cA:
        hand_soap = st.selectbox("Hand soap", ["(leave blank)", "Contractor", "Client"], index=0)
    with cB:
        paper_towels = st.selectbox("Paper towels", ["(leave blank)", "Contractor", "Client"], index=0)
    with cC:
        toilet_paper = st.selectbox("Toilet paper", ["(leave blank)", "Contractor", "Client"], index=0)

    st.subheader("Cover Page")
    include_cover = st.checkbox("Include cover page", value=True)
    use_standard_cover = st.checkbox("Use Torus standard cover letter", value=True)
    if use_standard_cover:
        cover_body = default_cover_letter(client)
        st.text_area("Cover letter (preview)", value=cover_body, height=220, disabled=True)
    else:
        cover_body = st.text_area("Cover letter body", value=st.session_state.get("cover_body_custom", ""), height=220)

    st.subheader("Contract Sections (On by default)")
    s1, s2, s3 = st.columns(3)
    with s1:
        include_employee_conduct = st.checkbox("Employee Conduct", value=True)
        include_on_site_storage = st.checkbox("On-Site Storage", value=True)
    with s2:
        include_compensation_section = st.checkbox("Compensation / Late Interest", value=True)
        include_modification = st.checkbox("Modification of Agreement", value=True)
    with s3:
        include_access = st.checkbox("Access", value=True)
        include_cancellation = st.checkbox("Cancellation", value=True)

    st.subheader("Payment (Optional)")
    st.caption("Compensation/Interest prints only if you enter a Compensation amount.")
    pay1, pay2, pay3, pay4 = st.columns([1, 1, 1, 1])
    with pay1:
        amount = st.text_input("Compensation amount (numbers only)", value="")
    with pay2:
        basis = st.selectbox("Basis", ["monthly", "annual", "per visit", "one-time clean"], index=0)
    with pay3:
        net_terms = st.selectbox("Net terms (days)", ["(leave blank)", "15", "30", "45", "60"], index=2)
    with pay4:
        late_interest = st.text_input("Late interest % (optional)", value="")

    st.subheader("Cleaning Plan & Notes")
    cleaning_plan = st.text_area("Cleaning Plan (optional)", height=120)
    notes = st.text_area("Notes", height=110)

    st.subheader("Cleaning Schedule")
    schedule_df = st.data_editor(
        st.session_state["schedule_df"],
        num_rows="dynamic",
        use_container_width=True,
        height=320,
    )

    st.subheader("RFP / PWS Analyzer (Optional)")
    uploads = st.file_uploader("Upload RFP/PWS", type=["pdf", "docx", "txt"], accept_multiple_files=True)

    a1, a2, a3 = st.columns([1, 1, 1])
    with a1:
        update_preview_btn = st.form_submit_button("Update Preview")
    with a2:
        analyze_btn = st.form_submit_button("Analyze with AI")
    with a3:
        generate_btn = st.form_submit_button("Generate Proposal")

# Persist edits
st.session_state["schedule_df"] = schedule_df
if not use_standard_cover:
    st.session_state["cover_body_custom"] = cover_body

# Normalize consumables
hand_soap_val = None if hand_soap == "(leave blank)" else hand_soap
paper_towels_val = None if paper_towels == "(leave blank)" else paper_towels
toilet_paper_val = None if toilet_paper == "(leave blank)" else toilet_paper

# Parse addresses
addresses = [x.strip() for x in (addresses_text or "").splitlines() if x.strip()]

# Convert schedule rows
schedule_rows = [
    (str(r.Task).strip(), bool(r.Daily), bool(r.Weekly), bool(r.Monthly))
    for r in st.session_state["schedule_df"].itertuples()
    if str(r.Task).strip()
]

def parse_float_or_none(x: str) -> Optional[float]:
    x = (x or "").strip()
    if not x:
        return None
    try:
        return float(x.replace(",", "").replace("$", ""))
    except Exception:
        return None

comp_amount = parse_float_or_none(amount)
late_interest_val = parse_float_or_none(late_interest)
net_terms_val = None if net_terms == "(leave blank)" else int(net_terms)

# Store preview inputs
if update_preview_btn or analyze_btn or generate_btn:
    st.session_state["last_inputs"] = {
        "client": client,
        "facility": facility,
        "begin": service_begin_date,
        "end": service_end_date,
        "days": int(days),
        "times": times,
        "addresses": addresses,
        "offices": int(offices),
        "conference": int(conference),
        "breaks": int(breaks),
        "baths": int(baths),
        "custom_rooms": st.session_state.get("custom_rooms", []),
        "consumables": {"hand_soap": hand_soap_val, "paper_towels": paper_towels_val, "toilet_paper": toilet_paper_val},
        "include_cover": bool(include_cover),
        "use_standard_cover": bool(use_standard_cover),
        "cover_body": cover_body,
        "cleaning_plan": cleaning_plan,
        "notes": notes,
        "payment": {"amount": comp_amount, "basis": basis, "net_terms": net_terms_val, "late_interest": late_interest_val},
        "sections": {
            "employee_conduct": bool(include_employee_conduct),
            "on_site_storage": bool(include_on_site_storage),
            "compensation": bool(include_compensation_section),
            "modification": bool(include_modification),
            "access": bool(include_access),
            "cancellation": bool(include_cancellation),
        },
        "schedule_rows": schedule_rows,
    }

# =========================
# PRINT PREVIEW (keeps same placement, just changes rendering)
# =========================
st.divider()
st.subheader("Preview")

li = st.session_state.get("last_inputs")
if not li:
    st.info("Fill out the form and press **Update Preview** to see the print preview.")
else:
    st.markdown(
        """
        <style>
          .page{
            background:#fff;
            max-width: 850px;
            margin: 0 auto;
            padding: 48px 56px;
            border: 1px solid #ddd;
            box-shadow: 0 2px 10px rgba(0,0,0,0.06);
            font-family: Arial, Helvetica, sans-serif;
            line-height: 1.35;
          }
          .doc-title{
            text-align:center;
            font-size: 20px;
            font-weight: 700;
            margin-bottom: 18px;
            letter-spacing: 0.5px;
          }
          h3{
            margin-top: 18px;
            margin-bottom: 8px;
            font-size: 14px;
            font-weight: 700;
          }
          ul{ margin-top: 6px; }
          .muted{ color:#666; font-style: italic; }
          table{
            width:100%;
            border-collapse: collapse;
            font-size: 12px;
          }
          th, td{
            border:1px solid #444;
            padding: 6px;
            vertical-align: top;
          }
          th{
            font-weight: 700;
            text-align:left;
          }
          .table-wrap{ margin-top: 8px; }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(build_print_preview_html(li), unsafe_allow_html=True)

# =========================
# AI ANALYSIS
# =========================
if analyze_btn:
    if not uploads:
        st.error("Please upload at least one RFP/PWS file.")
    else:
        try:
            full_text = "\n\n".join(extract_text(f) for f in uploads)
            if not full_text.strip():
                st.error("Could not extract text from the upload(s). If PDF is scanned, OCR is needed.")
            else:
                with st.spinner("Analyzing…"):
                    st.session_state["ai"] = analyze_rfp_with_ai(full_text)
                st.success("AI analysis complete.")
        except Exception as e:
            st.exception(e)

if st.session_state.get("ai"):
    ai = st.session_state["ai"]
    st.divider()
    st.subheader("AI Results")
    st.text_area("AI Cleaning Plan", ai.get("cleaning_plan_draft", ""), height=150)
    st.text_area("AI Scope of Work", ai.get("scope_of_work_draft", ""), height=150)

    qs = ai.get("clarifying_questions", [])
    if qs:
        st.write("**Clarifying questions**")
        for q in qs:
            st.write(f"- {q}")

    if st.button("Apply AI schedule to table"):
        rows = []
        for r in ai.get("schedule_rows", []):
            task = (r.get("task") or "").strip()
            if not task:
                continue
            rows.append((task, bool(r.get("daily")), bool(r.get("weekly")), bool(r.get("monthly"))))
        if rows:
            st.session_state["schedule_df"] = pd.DataFrame(rows, columns=["Task", "Daily", "Weekly", "Monthly"])
            st.success("Applied AI schedule. Scroll up—your schedule table is updated.")
            st.rerun()
        else:
            st.warning("AI did not return usable schedule rows.")

# =========================
# GENERATE DOC
# =========================
if generate_btn:
    p = ProposalInputs(
        client=client.strip(),
        facility_name=facility.strip(),
        service_begin_date=service_begin_date.strip(),
        service_end_date=service_end_date.strip(),
        service_addresses=addresses,
        days_per_week=int(days),
        cleaning_times=times.strip(),

        num_offices=int(offices),
        num_conference_rooms=int(conference),
        num_break_rooms=int(breaks),
        num_bathrooms=int(baths),

        custom_rooms=st.session_state.get("custom_rooms", []),

        hand_soap=hand_soap_val,
        paper_towels=paper_towels_val,
        toilet_paper=toilet_paper_val,

        include_cover_page=include_cover,
        cover_letter_body=cover_body,

        cleaning_plan=cleaning_plan,
        notes=notes,

        compensation_amount=comp_amount,
        compensation_basis=basis,
        net_terms_days=net_terms_val,
        late_interest_percent=late_interest_val,

        include_employee_conduct=include_employee_conduct,
        include_on_site_storage=include_on_site_storage,
        include_compensation_section=include_compensation_section,
        include_modification=include_modification,
        include_access=include_access,
        include_cancellation=include_cancellation,

        contractor_printed_name="Kary Jubilee",
        contractor_title="President, Torus Cleaning Services",
    )

    docx_bytes = build_doc(p, schedule_rows)

    st.success("Proposal generated.")
    st.download_button(
        "Download Word Proposal",
        data=docx_bytes,
        file_name=f"Torus_Cleaning_Agreement_{datetime.date.today().isoformat()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.download_button(
        "Download Inputs (JSON)",
        data=json.dumps(asdict(p), indent=2).encode("utf-8"),
        file_name=f"Torus_Inputs_{datetime.date.today().isoformat()}.json",
        mime="application/json",
    )
